import markdown
from docx import Document
import docx
from bs4 import BeautifulSoup
import os
from docx.shared import Pt
# conda activate pt
def markdown_to_html(markdown_text):
    """
    将Markdown文本转换为HTML
    """
    html = markdown.markdown(markdown_text)
    return html

def html_to_docx(html, output_file):
    """
    将HTML转换为Word文档
    """
    soup = BeautifulSoup(html, 'html.parser', from_encoding='utf-8')
    doc = Document()

    for element in soup.recursiveChildGenerator():
        name = getattr(element, 'name', None)
        if name is None:
            # 处理普通文本
            doc.add_paragraph(str(element))
        elif name == 'h1':
            # 处理标题1
            heading = doc.add_heading(level=0)
            heading.add_run(element.get_text()).font.size = Pt(24)
        elif name == 'h2':
            # 处理标题2
            heading = doc.add_heading(level=1)
            heading.add_run(element.get_text()).font.size = Pt(18)
        elif name == 'h3':
            # 处理标题3
            heading = doc.add_heading(level=2)
            heading.add_run(element.get_text()).font.size = Pt(14)
        elif name == 'p':
            # 处理段落
            paragraph = doc.add_paragraph()
            paragraph.add_run(element.get_text())
        elif name in ['ul', 'ol']:
            # 处理无序列表和有序列表
            for item in element.find_all('li', recursive=False):
                paragraph = doc.add_paragraph(style='ListBullet' if name == 'ul' else 'ListNumber')
                paragraph.add_run(item.get_text())
        elif name == 'a':
            # 处理链接
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(element.get_text())
            run.underline = True
            run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
        elif name == 'img':
            # 处理图片（需要图片的URL或本地路径）
            try:
                doc.add_picture(element['src'], width=docx.shared.Inches(4))
            except Exception as e:
                print(f"Error adding image: {e}")

    doc.save(output_file)

def markdown_file_to_docx(markdown_file_path, output_file):
    """
    从Markdown文件路径读取内容并转换为Word文档
    """
    if not os.path.exists(markdown_file_path):
        raise FileNotFoundError(f"Markdown文件未找到：{markdown_file_path}")

    with open(markdown_file_path, 'r', encoding='utf-8') as file:
        markdown_text = file.read()

    html = markdown_to_html(markdown_text)
    html_to_docx(html, output_file)
    print(f"Word文档已生成：{output_file}")

# 示例用法
markdown_file_path = "视觉盘点OCR方案打标V1-2025-02-12.md"  # 替换为你的Markdown文件路径
output_file = "output.docx"  # 输出的Word文档路径
markdown_file_to_docx(markdown_file_path, output_file)